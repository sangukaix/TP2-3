"""관광 전략기획서 Word 파일을 생성합니다.

LLM은 이 모듈에서 문장을 새로 만들지 않습니다. 이미 JSON Schema로 검증된
OpenAI 전략 보고서와 원자료 추세만 받아 문서 레이아웃으로 옮깁니다.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import textwrap
from typing import Any

import matplotlib

# FastAPI 서버에는 화면이 없으므로 차트 렌더러를 고정합니다.
matplotlib.use('Agg')

import matplotlib.pyplot as plt
import httpx
from matplotlib import font_manager
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = '203C60'
BLUE = '2E74B5'
MUTED = '6F8195'
PALE_BLUE = 'E8F2F8'
PALE_AQUA = 'EAF8F8'
GRID = 'D7E2EA'

# Windows 교육 환경의 기본 한글 폰트를 차트에 등록합니다.
# 없을 때는 matplotlib 기본 폰트로 안전하게 계속 생성합니다.
_KOREAN_FONT_PATH = Path('C:/Windows/Fonts/malgun.ttf')
if _KOREAN_FONT_PATH.exists():
    font_manager.fontManager.addfont(str(_KOREAN_FONT_PATH))
    plt.rcParams['font.family'] = font_manager.FontProperties(fname=str(_KOREAN_FONT_PATH)).get_name()
plt.rcParams['axes.unicode_minus'] = False


def _add_selected_official_image(document: Document, report: dict[str, Any], strategy: dict[str, Any]) -> None:
    """기획 Agent가 선택하고 출처가 보존된 공식 Open API 이미지만 최대 1장 넣습니다."""
    sources_by_id = {
        str(source.get('source_id')): source
        for source in report.get('evidence_sources') or []
        if source.get('source_id')
    }
    for source_id in strategy.get('visual_asset_source_ids') or []:
        source = sources_by_id.get(str(source_id)) or {}
        image_url = str(source.get('image_url') or '')
        if source.get('source_type') != 'open_api' or not image_url.startswith(('https://', 'http://')):
            continue
        try:
            response = httpx.get(image_url, timeout=12.0, follow_redirects=True)
            response.raise_for_status()
            if not response.headers.get('content-type', '').startswith('image/') or len(response.content) > 8_000_000:
                continue
            _add_heading(document, '공식 참고 이미지')
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(BytesIO(response.content), width=Inches(5.8))
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(
                caption.add_run(f"한국관광공사 Open API · {source.get('title', '관광자원')}"),
                size=8.5,
                color=MUTED,
            )
            return
        except (httpx.HTTPError, ValueError):
            continue


def _set_run_font(run: Any, *, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    """한글 Word에서도 폰트가 안정적으로 보이도록 동아시아 폰트를 같이 지정합니다."""
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Malgun Gothic')
    run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Malgun Gothic')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _shade(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    properties.append(shading)


def _cell_margins(cell: Any, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in('w:tcMar')
    if margins is None:
        margins = OxmlElement('w:tcMar')
        properties.append(margins)
    for side, value in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = margins.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            margins.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    """자동 너비에 맡기지 않고 Word 표의 실제 DXA 폭을 고정합니다."""
    table.autofit = False
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in('w:tblW')
    if table_width is None:
        table_width = OxmlElement('w:tblW')
        table_properties.append(table_width)
    table_width.set(qn('w:w'), str(sum(widths)))
    table_width.set(qn('w:type'), 'dxa')
    indent = table_properties.first_child_found_in('w:tblInd')
    if indent is None:
        indent = OxmlElement('w:tblInd')
        table_properties.append(indent)
    indent.set(qn('w:w'), '120')
    indent.set(qn('w:type'), 'dxa')
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            _cell_margins(cell)


def _add_text(cell: Any, text: str, *, bold: bool = False, color: str = INK, size: float = 9.5) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)


def _format_amount(value: int | float) -> str:
    return f'₩{round(float(value) / 100_000_000):,}억'


def _compact_text(value: Any, limit: int) -> str:
    """Word 표가 장문으로 밀리는 것을 막되, 생성된 기획안의 핵심 문장은 보존합니다."""
    text = ' '.join(str(value or '').split())
    if len(text) <= limit:
        return text
    shortened = text[:limit - 1].rsplit(' ', 1)[0]
    return f'{shortened}…'


def _execution_target(report: dict[str, Any]) -> tuple[float, float] | None:
    """사용자가 직접 입력한 목표율만 읽습니다. 입력이 없으면 임의 기본값을 만들지 않습니다."""
    scenario = report.get('execution_scenario') or {}
    if 'visitor_target_pct' not in scenario or 'spending_target_pct' not in scenario:
        return None
    visitor_pct = float(scenario['visitor_target_pct'])
    spending_pct = float(scenario['spending_target_pct'])
    return max(0, min(visitor_pct, 20)), max(0, min(spending_pct, 30))


def _ml_forecast_rows(report: dict[str, Any]) -> list[dict[str, Any]]:
    """저장 모델이 만든 자연추세 전망만 문서용 월별 행으로 정리합니다."""
    analysis = report.get('ml_analysis') or {}
    if analysis.get('status') != 'available':
        return []
    rows: list[dict[str, Any]] = []
    for forecast in (analysis.get('forecasts') or [])[:6]:
        month = str(forecast.get('month') or '')
        if len(month) == 6 and month.isdigit():
            month = f'{month[:4]}.{month[4:]}'
        visitors = forecast.get('visitors')
        spending = forecast.get('spending_krw')
        if visitors is None or spending is None:
            continue
        rows.append({'month': month, 'visitors': float(visitors), 'spending_krw': float(spending)})
    return rows


def _create_execution_comparison_chart(report: dict[str, Any]) -> BytesIO:
    """ML 자연추세와 사용자가 입력한 목표를 분리해 실제 단위로 표시합니다."""
    rows = _ml_forecast_rows(report)
    is_ml_forecast = bool(rows)
    if not rows:
        rows = [
            row for row in (report.get('monthly_trend') or [])[-6:]
            if row.get('visitors') is not None and row.get('spending_krw') is not None
        ]
    labels = [str(row.get('month') or '').replace('-', '.') for row in rows]
    natural_visitor_line = [float(row['visitors']) / 10_000 for row in rows]
    natural_spending_line = [float(row['spending_krw']) / 100_000_000 for row in rows]
    target = _execution_target(report)
    figure, axes = plt.subplots(2, 1, figsize=(7.05, 3.4), dpi=170, sharex=True)
    chart_specs = [
        (axes[0], natural_visitor_line, target[0] if target else None, '방문자 수(만 명)', '#24AFC0'),
        (axes[1], natural_spending_line, target[1] if target else None, '관광소비액(억 원)', '#8060C6'),
    ]
    for axis, natural, target_pct, ylabel, color in chart_specs:
        target_line = []
        if target_pct is not None and natural:
            last_index = max(1, len(natural) - 1)
            target_line = [value * (1 + target_pct / 100 * index / last_index) for index, value in enumerate(natural)]
        values = natural + target_line
        spread = max(values) - min(values)
        padding = max(spread * 0.3, max(values) * 0.006, 1)
        axis.set_ylim(max(0, min(values) - padding), max(values) + padding)
        natural_label = 'ML 자연추세' if is_ml_forecast else '최근 관측값'
        axis.plot(labels, natural, color=color, linewidth=2.4, marker='o', markersize=3, label=natural_label)
        if target_line:
            axis.plot(labels, target_line, color='#E86973', linewidth=1.8, linestyle='--', label='사용자 설정 목표')
        axis.set_ylabel(ylabel, color='#536B88', fontsize=7.5)
        axis.grid(axis='y', color='#DCE5EB', linewidth=0.6)
        axis.tick_params(axis='both', labelsize=6.8, colors='#536B88')
        axis.legend(loc='upper left', frameon=False, fontsize=6.8, ncol=2)
        for spine in axis.spines.values():
            spine.set_color('#C9D7E1')
    figure.tight_layout(pad=0.6)
    output = BytesIO()
    figure.savefig(output, format='png', transparent=False, facecolor='white')
    plt.close(figure)
    output.seek(0)
    return output


def _create_execution_timeline_chart(steps: list[dict[str, Any]]) -> BytesIO:
    """5단계 일정과 단계별 결과물을 한 줄의 프로젝트 타임라인으로 보여 줍니다."""
    visible_steps = (steps or [])[:5]
    positions = list(range(len(visible_steps)))
    figure, axis = plt.subplots(figsize=(7.05, 2.05), dpi=170)
    if visible_steps:
        axis.hlines(0, positions[0], positions[-1], color='#80B8C9', linewidth=2.2, zorder=1)
    axis.scatter(positions, [0] * len(positions), s=300, color='#3C8FA7', edgecolors='white', linewidths=2, zorder=2)
    for index, (position, step) in enumerate(zip(positions, visible_steps), start=1):
        axis.text(position, 0, str(index), ha='center', va='center', color='white', fontsize=8, fontweight='bold', zorder=3)
        axis.text(position, 0.28, _compact_text(step.get('schedule', ''), 18), ha='center', va='bottom', color='#405B77', fontsize=7.3, fontweight='bold')
        result = textwrap.fill(_compact_text(step.get('deliverable', ''), 24), width=9)
        axis.text(position, -0.3, result, ha='center', va='top', color='#6F8195', fontsize=6.8, linespacing=1.25)
    axis.set_xlim(-0.45, max(0.45, len(visible_steps) - 0.55))
    axis.set_ylim(-0.8, 0.72)
    axis.axis('off')
    figure.tight_layout(pad=0.4)
    output = BytesIO()
    figure.savefig(output, format='png', transparent=False, facecolor='white')
    plt.close(figure)
    output.seek(0)
    return output


def _create_trend_chart(trend: list[dict[str, Any]]) -> BytesIO:
    """원자료 월간 추세를 Word에 넣을 PNG로 만듭니다. LLM 생성 수치는 사용하지 않습니다."""
    months = [row['month'].replace('.', '\n') for row in trend]
    visitors = [row['visitors'] / 10_000 for row in trend]
    spending = [row['spending_krw'] / 100_000_000 for row in trend]
    figure, left_axis = plt.subplots(figsize=(7.05, 2.45), dpi=170)
    right_axis = left_axis.twinx()
    bars = left_axis.bar(months, visitors, color='#30BFD0', width=0.56, label='방문자 수')
    right_axis.plot(months, spending, color='#8560CC', linewidth=2.4, marker='o', markersize=3.6, label='관광소비액')
    left_axis.set_ylabel('방문자 수(만 명)', color='#536B88', fontsize=8)
    right_axis.set_ylabel('관광소비액(억 원)', color='#536B88', fontsize=8)
    left_axis.grid(axis='y', color='#DCE5EB', linewidth=0.6)
    left_axis.set_axisbelow(True)
    left_axis.tick_params(axis='both', labelsize=7, colors='#536B88')
    right_axis.tick_params(axis='y', labelsize=7, colors='#536B88')
    for spine in (*left_axis.spines.values(), *right_axis.spines.values()):
        spine.set_color('#C9D7E1')
    for bar, value in zip(bars, visitors):
        left_axis.text(bar.get_x() + bar.get_width() / 2, bar.get_height() * 0.5, f'{value:,.0f}', ha='center', va='center', color='white', fontsize=6.5)
    figure.tight_layout(pad=0.6)
    output = BytesIO()
    figure.savefig(output, format='png', transparent=False, facecolor='white')
    plt.close(figure)
    output.seek(0)
    return output


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    _set_run_font(run, size=14 if level == 1 else 11.5, bold=True, color=BLUE if level == 1 else INK)


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style='List Bullet')
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    _set_run_font(paragraph.add_run(text), size=10, color='3D526B')


def _add_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('TOUR INSIGHT | 원자료 기반 AI 전략기획서')
    _set_run_font(run, size=8, color=MUTED)


def _validate_document_structure(content: bytes) -> None:
    """서버에서 렌더러 없이도 잡을 수 있는 Word 구조 오류를 저장 전에 차단합니다.

    페이지 수·겹침은 별도 Word/LibreOffice 렌더 QA가 필요하지만, 필수 섹션 누락,
    과도한 본문, 5단계 표 손상, 예전의 미실행 정책효과 문구는 여기서 확정적으로 검사합니다.
    """
    checked = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in checked.paragraphs if paragraph.text.strip()]
    body_text = '\n'.join(paragraphs)
    table_text = '\n'.join(cell.text for table in checked.tables for row in table.rows for cell in row.cells)
    required_headings = [
        '1. 핵심 제안', '2. 최신 현황과 월간 변화', '3. 문제 / 제안과 판단 근거',
        '4. 해결 방법', '5. 집행 방법', '6. ML 자연추세와 사업 목표',
        '7. 참고한 공식 사례·데이터',
    ]
    missing = [heading for heading in required_headings if heading not in paragraphs]
    if missing:
        raise ValueError(f"Word 필수 섹션이 누락되었습니다: {', '.join(missing)}")
    if len(body_text) + len(table_text) > 14_000:
        raise ValueError('Word 본문이 문서 검토 상한을 넘었습니다. 긴 문장을 더 요약해 주세요.')
    if '미실행 (최근 월)' in table_text or '추가 관광소비' in table_text:
        raise ValueError('정책 인과효과로 오해할 수 있는 이전 비교 표기가 남아 있습니다.')
    execution_tables = [
        table for table in checked.tables
        if table.rows and [cell.text.strip() for cell in table.rows[0].cells] == ['순서', '기간', '실행 작업', '완료 산출물']
    ]
    if len(execution_tables) != 1 or len(execution_tables[0].rows) != 6:
        raise ValueError('집행 방법은 머리글과 정확히 5개 실행 단계로 구성해야 합니다.')


def create_strategy_proposal_document(report: dict[str, Any]) -> BytesIO:
    """최대 5쪽 안에서 빠르게 검토할 수 있는 도표 중심 Word 기획서를 생성합니다."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.66)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)
    _add_footer(document)

    normal_style = document.styles['Normal']
    normal_style.paragraph_format.space_after = Pt(5)
    normal_style.paragraph_format.line_spacing = 1.1
    # Style의 Font 객체는 Run과 구조가 달라 별도 설정합니다.
    normal_style.font.name = '맑은 고딕'
    normal_style._element.rPr.rFonts.set(qn('w:ascii'), 'Malgun Gothic')
    normal_style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Malgun Gothic')
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor.from_string('3D526B')

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run(f"{report['region_name']} 관광 기획안")
    _set_run_font(run, size=22, bold=True, color=INK)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    _set_run_font(subtitle.add_run(f"분석 기간 {report['period']} | 관광데이터랩 원자료와 공식 관광자료 기반"), size=9.5, color=MUTED)
    _add_heading(document, '1. 핵심 제안')
    if report.get('planning_brief'):
        from .planning_brief import brief_summary
        conditions = document.add_paragraph(brief_summary(report['planning_brief']))
        for item in conditions.runs:
            _set_run_font(item, size=9, color=MUTED)
    callout = document.add_table(rows=1, cols=1)
    _set_table_geometry(callout, [9360])
    _shade(callout.cell(0, 0), PALE_AQUA)
    _add_text(callout.cell(0, 0), _compact_text(report['summary'], 300), size=10.5, color=INK)

    _add_heading(document, '2. 최신 현황과 월간 변화')
    findings = report['observed_findings'][:3]
    metric_table = document.add_table(rows=1, cols=2)
    _set_table_geometry(metric_table, [3000, 6360])
    _shade(metric_table.cell(0, 0), PALE_BLUE)
    _shade(metric_table.cell(0, 1), PALE_BLUE)
    _add_text(metric_table.cell(0, 0), '관측 지표', bold=True, size=9)
    _add_text(metric_table.cell(0, 1), '최근 값 및 해석', bold=True, size=9)
    for finding in findings:
        cells = metric_table.add_row().cells
        _add_text(cells[0], finding['metric'], bold=True, size=9)
        _add_text(cells[1], f"{finding['value']} | {_compact_text(finding['interpretation'], 84)}", size=9)
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    document.add_picture(_create_trend_chart(report['monthly_trend']), width=Inches(6.45))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(caption.add_run('그림 1. 월간 순 방문자 수와 외지인 관광소비액 추이'), size=8.5, color=MUTED)

    strategy = sorted(report['strategies'], key=lambda item: item['priority'])[0]
    _add_selected_official_image(document, report, strategy)

    _add_heading(document, '3. 문제 / 제안과 판단 근거')
    diagnosis_table = document.add_table(rows=2, cols=2)
    _set_table_geometry(diagnosis_table, [1760, 7600])
    diagnosis_rows = [
        ('문제 / 제안', _compact_text(strategy['problem_to_solve'], 220)),
        ('이렇게 판단한 이유', _compact_text(strategy['comparison_analysis'], 300)),
    ]
    for row, (label, value) in zip(diagnosis_table.rows, diagnosis_rows):
        _shade(row.cells[0], PALE_BLUE)
        _add_text(row.cells[0], label, bold=True, size=9)
        _add_text(row.cells[1], value, size=9)

    _add_heading(document, '4. 해결 방법')
    solution_table = document.add_table(rows=2, cols=2)
    _set_table_geometry(solution_table, [1760, 7600])
    solution_rows = [('실행 방향', _compact_text(strategy['solution'], 240)), ('예산 준비', _compact_text(strategy['budget'], 180))]
    for row, (label, value) in zip(solution_table.rows, solution_rows):
        _shade(row.cells[0], PALE_AQUA)
        _add_text(row.cells[0], label, bold=True, size=9)
        _add_text(row.cells[1], value, size=9)

    _add_heading(document, '5. 집행 방법')
    execution_table = document.add_table(rows=1, cols=4)
    _set_table_geometry(execution_table, [700, 1300, 4560, 2800])
    for cell, header in zip(execution_table.rows[0].cells, ['순서', '기간', '실행 작업', '완료 산출물']):
        _shade(cell, PALE_BLUE)
        _add_text(cell, header, bold=True, size=8.5)
    for index, step in enumerate(strategy.get('implementation_steps') or [], start=1):
        cells = execution_table.add_row().cells
        _add_text(cells[0], str(step.get('step', index)), bold=True, color=BLUE, size=8.5)
        _add_text(cells[1], step.get('schedule', ''), bold=True, size=8.5)
        _add_text(cells[2], _compact_text(step.get('task', ''), 90), size=8.5)
        _add_text(cells[3], _compact_text(step.get('deliverable', ''), 48), size=8.5)
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    document.add_picture(_create_execution_timeline_chart(strategy.get('implementation_steps') or []), width=Inches(6.45))
    timeline_caption = document.add_paragraph()
    timeline_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(timeline_caption.add_run('그림 2. 5단계 실행 일정과 단계별 결과물'), size=8.3, color=MUTED)

    _add_heading(document, '6. ML 자연추세와 사업 목표')
    forecasts = _ml_forecast_rows(report)
    target = _execution_target(report)
    if forecasts:
        last_forecast = forecasts[-1]
        if target:
            visitor_pct, spending_pct = target
            target_visitors = round(last_forecast['visitors'] * (1 + visitor_pct / 100))
            target_spending = round(last_forecast['spending_krw'] * (1 + spending_pct / 100))
            comparison_table = document.add_table(rows=2, cols=3)
            _set_table_geometry(comparison_table, [3120, 3120, 3120])
            headers = ['ML 자연추세', '사용자 설정 목표', '추가로 필요한 수준']
            for cell, header in zip(comparison_table.rows[0].cells, headers):
                _shade(cell, PALE_AQUA)
                _add_text(cell, header, bold=True, size=8.5)
            _add_text(comparison_table.cell(1, 0), f"{last_forecast['month']}\n방문 {last_forecast['visitors']:,.0f}명\n소비 {_format_amount(last_forecast['spending_krw'])}", size=8.8)
            _add_text(comparison_table.cell(1, 1), f"방문 {target_visitors:,}명 (+{visitor_pct:g}%)\n소비 {_format_amount(target_spending)} (+{spending_pct:g}%)", size=8.8)
            _add_text(comparison_table.cell(1, 2), f"방문 {target_visitors - round(last_forecast['visitors']):,}명\n소비 {_format_amount(target_spending - last_forecast['spending_krw'])}", bold=True, color=BLUE, size=8.8)
        else:
            comparison_table = document.add_table(rows=2, cols=2)
            _set_table_geometry(comparison_table, [3300, 6060])
            for cell, header in zip(comparison_table.rows[0].cells, ['전망 마지막 달', '저장 모델의 자연추세 전망']):
                _shade(cell, PALE_AQUA)
                _add_text(cell, header, bold=True, size=8.5)
            _add_text(comparison_table.cell(1, 0), last_forecast['month'], bold=True, size=9)
            _add_text(comparison_table.cell(1, 1), f"방문 {last_forecast['visitors']:,.0f}명 · 관광소비 {_format_amount(last_forecast['spending_krw'])}", size=9)
        document.add_paragraph().paragraph_format.space_after = Pt(1)
        document.add_picture(_create_execution_comparison_chart(report), width=Inches(6.45))
        comparison_note = document.add_paragraph()
        comparison_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note = '그림 3. 저장 ML 모델의 월별 자연추세'
        if target:
            note += '와 사용자 설정 목표'
        _set_run_font(comparison_note.add_run(note), size=8.3, color=MUTED)
        forecast_caution = document.add_paragraph()
        forecast_caution.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(
            forecast_caution.add_run('자연추세는 과거 이력 기반 전망이며, 사업의 인과효과나 추가 매출 예측이 아닙니다.'),
            size=8.2,
            color=MUTED,
        )
    else:
        no_forecast = document.add_paragraph('검증된 저장 모델 전망이 없어 임의의 미실행·실행 수치를 만들지 않았습니다.')
        for run in no_forecast.runs:
            _set_run_font(run, size=9, color=MUTED)
    effect_callout = document.add_table(rows=1, cols=1)
    _set_table_geometry(effect_callout, [9360])
    _shade(effect_callout.cell(0, 0), PALE_AQUA)
    _add_text(effect_callout.cell(0, 0), f"기대할 수 있는 변화 · {_compact_text(strategy['expected_effect'], 180)}", size=9.2, color=INK)

    _add_heading(document, '7. 참고한 공식 사례·데이터')
    all_sources = report.get('evidence_sources') or []
    benchmark_sources = [source for source in all_sources if source.get('source_type') == 'benchmark_case'][:3]
    if benchmark_sources:
        case_table = document.add_table(rows=1, cols=2)
        _set_table_geometry(case_table, [3100, 6260])
        for cell, header in zip(case_table.rows[0].cells, ['공식 사례', '전략에 참고한 내용']):
            _shade(cell, PALE_AQUA)
            _add_text(cell, header, bold=True, size=8.2)
        for source in benchmark_sources:
            cells = case_table.add_row().cells
            _add_text(cells[0], _compact_text(source.get('title', ''), 55), bold=True, size=7.9)
            _add_text(cells[1], _compact_text(source.get('summary', ''), 115), size=7.9)

    source_type_names = {
        'dataset': '관광데이터랩',
        'open_api': '관광 Open API',
        'official_web': '공식 문서',
        'rag': '공식 자료',
        'benchmark_case': '공식 성공사례',
    }
    supporting_sources = [source for source in all_sources if source.get('source_type') != 'benchmark_case'][:5]
    for source in supporting_sources:
        source_line = document.add_paragraph(style='List Bullet')
        source_line.paragraph_format.left_indent = Inches(0.15)
        source_line.paragraph_format.space_after = Pt(1)
        _set_run_font(
            source_line.add_run(f"{source_type_names.get(source.get('source_type'), '공식 자료')} · {_compact_text(source.get('title', ''), 95)}"),
            size=8.2,
            color=MUTED,
        )

    output = BytesIO()
    document.save(output)
    content = output.getvalue()
    _validate_document_structure(content)
    return BytesIO(content)
