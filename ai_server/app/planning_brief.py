"""사업 여건 계약. 사용자 입력은 관측 사실/공용 RAG와 분리해 보관합니다."""
from __future__ import annotations

from datetime import date
from hashlib import sha256
from io import BytesIO
import json
from pathlib import Path
from typing import Literal
from zipfile import ZipFile, BadZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from pydantic import BaseModel, ConfigDict, Field, model_validator


class BriefReference(BaseModel):
    model_config = ConfigDict(extra='forbid', str_strip_whitespace=True)
    name: str = Field(min_length=1, max_length=150)
    text: str = Field(min_length=1, max_length=6000)


class PlanningBrief(BaseModel):
    """정해지지 않은 값은 null/unknown으로 유지합니다. 0원이나 확정으로 바꾸지 않습니다."""
    model_config = ConfigDict(extra='forbid', str_strip_whitespace=True)
    version: Literal[1] = 1
    region_code: str = Field(pattern=r'^\d{2,5}$')
    budget_status: Literal['unknown', 'indicative', 'confirmed'] = 'unknown'
    budget_min_krw: int | None = Field(default=None, ge=1, le=1_000_000_000_000, strict=True)
    budget_max_krw: int | None = Field(default=None, ge=1, le=1_000_000_000_000, strict=True)
    budget_hard_limit: bool = False
    schedule_status: Literal['unknown', 'flexible', 'fixed'] = 'unknown'
    start_date: date | None = None
    end_date: date | None = None
    # 빈 입력과 ‘알 수 없음’을 구분해 AI가 시설·인력이 없다고 단정하지 않게 합니다.
    resources_status: Literal['unknown', 'known'] = 'unknown'
    resources_confirmed: str = Field(default='', max_length=1500)
    resources_possible: str = Field(default='', max_length=1500)
    constraints_status: Literal['unknown', 'known'] = 'unknown'
    hard_constraints: str = Field(default='', max_length=2000)
    preferences: str = Field(default='', max_length=1000)
    field_context: str = Field(default='', max_length=2500)
    references: list[BriefReference] = Field(default_factory=list, max_length=3)

    @model_validator(mode='before')
    @classmethod
    def infer_legacy_input_status(cls, value):
        """상태 필드가 없던 기존 브라우저 초안도 내용을 잃지 않고 읽습니다."""
        if not isinstance(value, dict):
            return value
        data = dict(value)
        if 'resources_status' not in data:
            data['resources_status'] = 'known' if data.get('resources_confirmed') or data.get('resources_possible') else 'unknown'
        if 'constraints_status' not in data:
            data['constraints_status'] = 'known' if data.get('hard_constraints') else 'unknown'
        return data

    @model_validator(mode='after')
    def check_conditions(self):
        if self.budget_status == 'unknown':
            if self.budget_min_krw is not None or self.budget_max_krw is not None or self.budget_hard_limit:
                raise ValueError('예산 미정 상태에서는 금액과 상한을 지정할 수 없습니다.')
        elif self.budget_max_krw is None:
            raise ValueError('예산 금액 또는 범위의 최대 금액을 입력해 주세요.')
        if self.budget_min_krw and self.budget_max_krw and self.budget_min_krw > self.budget_max_krw:
            raise ValueError('최소 예산은 최대 예산보다 클 수 없습니다.')
        if self.schedule_status == 'unknown' and (self.start_date or self.end_date):
            raise ValueError('일정 미정 상태에서는 날짜를 지정할 수 없습니다.')
        if self.schedule_status != 'unknown' and not (self.start_date and self.end_date):
            raise ValueError('사업 시작일과 종료일을 모두 입력해 주세요.')
        if self.start_date and self.end_date and self.start_date > self.end_date:
            raise ValueError('사업 종료일은 시작일보다 빠를 수 없습니다.')
        if self.resources_status == 'unknown' and (self.resources_confirmed or self.resources_possible):
            raise ValueError('시설·인력 미정 상태에서는 내용을 입력할 수 없습니다.')
        if self.constraints_status == 'unknown' and self.hard_constraints:
            raise ValueError('필수 조건 미정 상태에서는 내용을 입력할 수 없습니다.')
        return self


def brief_fingerprint(brief: dict | None) -> str:
    if not brief:
        return ''
    return sha256(json.dumps(brief, ensure_ascii=False, sort_keys=True).encode()).hexdigest()


def without_reference_text(brief: PlanningBrief | None) -> PlanningBrief | None:
    """생성 뒤 보관·응답되는 조건에서는 첨부 본문을 제거합니다."""
    return brief.model_copy(update={'references': []}) if brief else None


def brief_summary(brief: dict | None) -> str:
    """화면/다운로드용 짧은 조건 요약. 출처 수치가 아닌 사용자 설정임을 명시합니다."""
    if not brief:
        return ''
    amount = brief.get('budget_max_krw')
    budget = f"{amount:,}원" if amount else '미정'
    budget += ' (상한)' if brief.get('budget_hard_limit') else ''
    period = f"{brief['start_date']} ~ {brief['end_date']}" if brief.get('start_date') and brief.get('end_date') else '미정'
    return f"사용자 입력 조건 | 예산 {budget} · 일정 {period}"


def extract_brief_reference(filename: str, content: bytes) -> dict:
    """첨부 문서의 필요한 텍스트만 메모리에서 읽고, 원본은 저장하지 않습니다."""
    if not content or len(content) > 2_000_000:
        raise ValueError('파일은 2MB 이하로 첨부해 주세요.')
    name = Path(filename.replace('\\', '/')).name[:150]
    suffix = Path(name).suffix.lower()
    if suffix in ('.txt', '.md'):
        try:
            text = content.decode('utf-8-sig')
        except UnicodeDecodeError:
            text = content.decode('cp949')
    elif suffix in ('.docx', '.hwpx', '.xlsx'):
        try:
            with ZipFile(BytesIO(content)) as archive:
                files = archive.infolist()
                if len(files) > 500 or sum(item.file_size for item in files) > 12_000_000:
                    raise ValueError('압축 해제 크기가 너무 큰 문서는 첨부할 수 없습니다.')
            if suffix == '.docx':
                document = Document(BytesIO(content))
                text = '\n'.join([p.text for p in document.paragraphs] + [
                    ' | '.join(c.text for c in row.cells) for t in document.tables for row in t.rows
                ])
            elif suffix == '.hwpx':
                text = _extract_hwpx_text(BytesIO(content))
            else:
                text = _extract_xlsx_text(BytesIO(content))
        except (BadZipFile, KeyError) as exc:
            raise ValueError('올바른 Word·한글(HWPX)·Excel 파일을 첨부해 주세요.') from exc
    elif suffix == '.pdf':
        try:
            reader = PdfReader(BytesIO(content))
            if reader.is_encrypted:
                raise ValueError('암호가 설정된 PDF는 텍스트를 읽을 수 없습니다.')
            text = '\n'.join(page.extract_text() or '' for page in reader.pages[:40])
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError('텍스트를 읽을 수 있는 PDF를 첨부해 주세요.') from exc
    else:
        raise ValueError('한글(HWPX), Word(DOCX), PDF, TXT, Excel(XLSX) 파일을 첨부해 주세요. HWP는 HWPX 또는 PDF로 저장해 주세요.')
    text = text.replace('\x00', '').strip()
    if not text:
        raise ValueError('문서에서 텍스트를 찾지 못했습니다.')
    if len(text) > 6000:
        raise ValueError('문서 내용이 6,000자를 넘습니다. 필요한 부분만 별도 문서로 첨부해 주세요.')
    return BriefReference(name=name, text=text).model_dump()


def _extract_hwpx_text(buffer: BytesIO) -> str:
    """HWPX는 ZIP 안 XML 문서라 외부 변환기 없이 본문 텍스트만 읽을 수 있습니다."""
    import xml.etree.ElementTree as etree

    with ZipFile(buffer) as archive:
        sections = sorted(name for name in archive.namelist() if name.startswith('Contents/section') and name.endswith('.xml'))
        if not sections:
            raise ValueError('본문이 없는 HWPX 파일입니다.')
        return '\n'.join(''.join(element.itertext()) for name in sections for element in [etree.fromstring(archive.read(name))])


def _extract_xlsx_text(buffer: BytesIO) -> str:
    """Excel은 첫 10개 시트·각 500행만 읽어, 참고자료 처리 시간을 제한합니다."""
    workbook = load_workbook(buffer, read_only=True, data_only=True)
    rows: list[str] = []
    try:
        for worksheet in workbook.worksheets[:10]:
            rows.append(f'[{worksheet.title}]')
            for row in worksheet.iter_rows(max_row=500, values_only=True):
                values = [str(value).strip() for value in row[:50] if value is not None and str(value).strip()]
                if values:
                    rows.append(' | '.join(values))
    finally:
        workbook.close()
    return '\n'.join(rows)
